import os
import requests
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document
from langchain.document_loaders import WebBaseLoader
from tqdm.notebook import tqdm_notebook
import re

# ---------------------------
# Download helper
# ---------------------------
def download_file_with_user_agent(url, local_filename):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    try:
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            with open(local_filename, 'wb') as f:
                f.write(response.content)
            return True
        else:
            print(f"Failed to retrieve {url}, status code: {response.status_code}")
            return False
    except Exception as e:
        print(f"Error downloading {url}: {e}")
        return False

# ---------------------------
# Load a SINGLE document
# ---------------------------
def load_document(url):
    """
    Loads a single document from a URL (PDF, DOCX, or webpage).
    Returns extracted text as a single string.
    """
    text = ""

    if url.endswith(".pdf"):
        local_filename = "temp_downloaded_file.pdf"
        try:
            response = requests.get(url)
            if response.status_code == 200:
                pdf_reader = PdfReader(BytesIO(response.content))
                text = "\n".join(page.extract_text() or "" for page in pdf_reader.pages)
            else:
                if download_file_with_user_agent(url, local_filename):
                    with open(local_filename, "rb") as f:
                        pdf_reader = PdfReader(f)
                        text = "\n".join(page.extract_text() or "" for page in pdf_reader.pages)
                    os.remove(local_filename)
        except Exception as e:
            print(f"Error loading PDF from {url}: {e}")

    elif url.endswith(".docx"):
        local_filename = "temp_downloaded_file.docx"
        try:
            response = requests.get(url)
            if response.status_code == 200:
                document = Document(BytesIO(response.content))
                text = "\n".join(p.text for p in document.paragraphs)
            else:
                if download_file_with_user_agent(url, local_filename):
                    document = Document(local_filename)
                    text = "\n".join(p.text for p in document.paragraphs)
                    os.remove(local_filename)
        except Exception as e:
            print(f"Error loading DOCX from {url}: {e}")

    else:
        try:
            loader = WebBaseLoader(url)
            web_texts = loader.load()
            text = " ".join(item.page_content for item in web_texts)
        except Exception as e:
            print(f"Error loading webpage from {url}: {e}")

    return text.strip()

# ---------------------------
# Chunking
# ---------------------------
def chunk_text(text, max_chunk_size=500, overlap_size=50):
    """
    Splits a single text into overlapping word chunks.
    """
    def count_words(txt):
        return len(txt.split())

    def split_into_sentences(txt):
        return re.split(r'(?<=[.!?])\s+(?=[A-Z])', txt)

    paragraphs = re.split(r"\n\s*\n", text.strip())
    paragraphs = [p.strip() for p in paragraphs if p.strip()]

    chunks = []
    current_chunk = []
    current_word_count = 0

    for paragraph in paragraphs:
        paragraph_word_count = count_words(paragraph)

        if paragraph_word_count > max_chunk_size:
            if current_chunk:
                chunks.append(" ".join(current_chunk))
                current_chunk, current_word_count = [], 0

            sentences = split_into_sentences(paragraph)
            sentence_chunk, sentence_word_count = [], 0

            for sentence in sentences:
                sentence_words = count_words(sentence)
                if sentence_word_count + sentence_words > max_chunk_size:
                    chunks.append(" ".join(sentence_chunk))
                    sentence_chunk = sentence_chunk[-overlap_size//10:] if overlap_size else []
                    sentence_word_count = count_words(" ".join(sentence_chunk))
                sentence_chunk.append(sentence)
                sentence_word_count += sentence_words

            if sentence_chunk:
                chunks.append(" ".join(sentence_chunk))

        elif current_word_count + paragraph_word_count > max_chunk_size:
            chunks.append(" ".join(current_chunk))
            current_chunk = current_chunk[-overlap_size//10:] if overlap_size else []
            current_word_count = count_words(" ".join(current_chunk))
            current_chunk.append(paragraph)
            current_word_count += paragraph_word_count

        else:
            current_chunk.append(paragraph)
            current_word_count += paragraph_word_count

    if current_chunk:
        chunks.append(" ".join(current_chunk))

    return chunks

# ---------------------------
# Unified entry point
# ---------------------------
def get_docs(url):
    """
    Input: single URL string
    Output: list of text chunks (list[str])
    """
    raw_text = load_document(url)
    if not raw_text:
        return []
    return chunk_text(raw_text)
