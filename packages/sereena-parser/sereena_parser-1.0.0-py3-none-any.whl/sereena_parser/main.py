import os
import re
import json
import logging
import threading
import argparse
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, asdict
from pathlib import Path
import time

# Third-party imports
try:
    import pypdf
    import pdfplumber
    from docx import Document
    import spacy
    from spacy.matcher import Matcher
    import pandas as pd
    from email_validator import validate_email, EmailNotValidError
except ImportError as e:
    print(f"Missing required dependency: {e}")
    print("Please install with: pip install sereena-parser")
    sys.exit(1)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('sereena_parser.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

@dataclass
class PersonalInfo:
    """Data class for personal information"""
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None

@dataclass
class Education:
    """Data class for education information"""
    degree: Optional[str] = None
    institution: Optional[str] = None
    year: Optional[str] = None
    gpa: Optional[str] = None

@dataclass
class Experience:
    """Data class for work experience"""
    job_title: Optional[str] = None
    company: Optional[str] = None
    duration: Optional[str] = None
    description: Optional[str] = None

@dataclass
class ResumeData:
    """Main data class for parsed resume information"""
    personal_info: PersonalInfo
    skills: List[str]
    education: List[Education]
    experience: List[Experience]
    certifications: List[str]
    languages: List[str]
    summary: Optional[str] = None
    raw_text: Optional[str] = None

class TextExtractor:
    """Handles text extraction from different file formats"""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF using multiple methods for better accuracy"""
        text = ""
        
        try:
            # Method 1: pdfplumber (better for formatted PDFs)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # If pdfplumber fails or returns empty, try pypdf
            if not text.strip():
                with open(file_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        
        except Exception as e:
            logger.error(f"Error extracting PDF {file_path}: {str(e)}")
            raise
            
        return text.strip()
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ""
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text += cell.text.strip() + " "
                    if row_text.strip():
                        text += row_text + "\n"
                    
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting DOCX {file_path}: {str(e)}")
            raise

class EntityExtractor:
    """Handles entity extraction using NLP and pattern matching"""
    
    def __init__(self):
        """Initialize NLP models and patterns"""
        try:
            # Load spaCy model
            self.nlp = spacy.load("en_core_web_sm")
            logger.info("Loaded spaCy model successfully")
        except OSError:
            logger.error("spaCy model not found. Install with: python -m spacy download en_core_web_sm")
            raise
        
        self.matcher = Matcher(self.nlp.vocab)
        self._setup_patterns()
        self._load_skill_database()
    
    def _setup_patterns(self):
        """Setup regex patterns for entity extraction"""
        self.patterns = {
            'email': re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'),
            'phone': re.compile(r'(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'),
            'gpa': re.compile(r'(?:GPA|gpa)\s*[:\-]?\s*([0-4]\.\d+|[0-4]/4|[0-9]{1,3}%)', re.IGNORECASE),
            'year': re.compile(r'\b(19|20)\d{2}\b'),
            'degree': re.compile(r'\b(?:Bachelor|Master|PhD|Doctorate|Associates?|BS|BA|MS|MA|MBA|PhD)\b.*?(?:in|of)?\s+([^,\n\.]+)', re.IGNORECASE),
            'experience_years': re.compile(r'(\d+)[\+]?\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)', re.IGNORECASE),
        }
    
    def _load_skill_database(self):
        """Load comprehensive skills database"""
        self.skills_db = {
            'programming': [
                'Python', 'Java', 'JavaScript', 'C++', 'C#', 'Ruby', 'Go', 'Rust', 'Swift',
                'Kotlin', 'PHP', 'R', 'MATLAB', 'Scala', 'Perl', 'TypeScript', 'Dart',
                'C', 'Objective-C', 'VB.NET', 'F#', 'Haskell', 'Clojure', 'Elixir'
            ],
            'web_technologies': [
                'HTML', 'CSS', 'React', 'Angular', 'Vue.js', 'Node.js', 'Django', 'Flask',
                'Express.js', 'Spring Boot', 'ASP.NET', 'Laravel', 'Ruby on Rails',
                'jQuery', 'Bootstrap', 'Sass', 'Less', 'Webpack', 'Gulp', 'Next.js',
                'Nuxt.js', 'Svelte', 'GraphQL', 'REST API'
            ],
            'databases': [
                'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Oracle', 'SQL Server',
                'SQLite', 'Cassandra', 'DynamoDB', 'Neo4j', 'CouchDB', 'InfluxDB',
                'MariaDB', 'Firebase', 'Elasticsearch'
            ],
            'cloud_platforms': [
                'AWS', 'Azure', 'Google Cloud', 'GCP', 'Heroku', 'DigitalOcean',
                'IBM Cloud', 'Oracle Cloud', 'Alibaba Cloud', 'Linode', 'Vultr'
            ],
            'tools': [
                'Git', 'Docker', 'Kubernetes', 'Jenkins', 'Ansible', 'Terraform',
                'Jira', 'Confluence', 'Slack', 'Teams', 'Trello', 'Asana',
                'GitLab', 'GitHub', 'Bitbucket', 'CircleCI', 'Travis CI'
            ],
            'data_science': [
                'Machine Learning', 'Deep Learning', 'TensorFlow', 'PyTorch', 'Scikit-learn',
                'Pandas', 'NumPy', 'Matplotlib', 'Seaborn', 'Tableau', 'Power BI',
                'Jupyter', 'Apache Spark', 'Hadoop', 'Kafka', 'Airflow', 'MLflow'
            ],
            'mobile': [
                'Android', 'iOS', 'React Native', 'Flutter', 'Xamarin', 'Ionic',
                'Swift', 'Objective-C', 'Kotlin', 'Java'
            ],
            'devops': [
                'Docker', 'Kubernetes', 'Jenkins', 'Ansible', 'Terraform', 'Puppet',
                'Chef', 'Vagrant', 'Prometheus', 'Grafana', 'ELK Stack', 'Nagios'
            ]
        }
        
        # Flatten skills for easier matching
        self.all_skills = []
        for category in self.skills_db.values():
            self.all_skills.extend(category)
        
        # Remove duplicates while preserving order
        self.all_skills = list(dict.fromkeys(self.all_skills))
    
    def extract_personal_info(self, text: str) -> PersonalInfo:
        """Extract personal information from text"""
        personal_info = PersonalInfo()
        
        # Extract email
        email_matches = self.patterns['email'].findall(text)
        for email in email_matches:
            try:
                validated_email = validate_email(email)
                personal_info.email = validated_email.email
                break  # Take the first valid email
            except EmailNotValidError:
                continue
        
        # Extract phone
        phone_match = self.patterns['phone'].search(text)
        if phone_match:
            personal_info.phone = phone_match.group()
        
        # Extract name using NER
        doc = self.nlp(text[:800])  # Process first 800 chars for better performance
        for ent in doc.ents:
            if ent.label_ == "PERSON" and not personal_info.name:
                # Filter out common false positives
                name = ent.text.strip()
                if len(name.split()) >= 2 and not any(word in name.lower() for word in ['university', 'college', 'company']):
                    personal_info.name = name
                    break
        
        # Extract location
        locations = []
        for ent in doc.ents:
            if ent.label_ in ["GPE", "LOC"] and len(ent.text.strip()) > 2:
                locations.append(ent.text.strip())
        
        if locations:
            # Prefer locations that appear in typical location contexts
            location_keywords = ['based in', 'located in', 'from', 'living in', 'residing in']
            for keyword in location_keywords:
                for location in locations:
                    if keyword + ' ' + location.lower() in text.lower():
                        personal_info.location = location
                        break
                if personal_info.location:
                    break
            
            if not personal_info.location:
                personal_info.location = locations[0]
        
        return personal_info
    
    def extract_skills(self, text: str) -> List[str]:
        """Extract skills from text using pattern matching and NLP"""
        found_skills = set()
        text_lower = text.lower()
        
        # Direct skill matching
        for skill in self.all_skills:
            skill_lower = skill.lower()
            # Use word boundaries to avoid partial matches
            if re.search(r'\b' + re.escape(skill_lower) + r'\b', text_lower):
                found_skills.add(skill)
        
        # Pattern-based skill extraction
        skill_patterns = [
            r'(?:skilled?\s+(?:in|with)|experience\s+(?:in|with)|proficient\s+(?:in|with)|knowledge\s+of|familiar\s+with)\s+([^.,\n]{3,50})',
            r'(?:technologies|skills|tools|languages|frameworks|platforms):\s*([^.\n]{10,200})',
            r'(?:technical\s+skills|core\s+competencies|expertise):\s*([^.\n]{10,200})',
        ]
        
        for pattern in skill_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                skills_text = match.group(1).lower()
                for skill in self.all_skills:
                    if skill.lower() in skills_text:
                        found_skills.add(skill)
        
        return sorted(list(found_skills))
    
    def extract_education(self, text: str) -> List[Education]:
        """Extract education information"""
        educations = []
        
        # Split text into sections
        education_sections = re.split(r'\b(?:education|academic|qualifications|degrees?)\b', text, flags=re.IGNORECASE)
        
        if len(education_sections) > 1:
            # Take text after education header, limit to reasonable size
            education_text = education_sections[1][:1500]
        else:
            # Look for degree patterns in entire text
            education_text = text
        
        # Find degree patterns
        degree_matches = list(self.patterns['degree'].finditer(education_text))
        
        # Use NLP to find organizations (potential schools)
        doc = self.nlp(education_text)
        organizations = [ent.text for ent in doc.ents if ent.label_ == "ORG" and len(ent.text) > 2]
        
        # Find years
        years = self.patterns['year'].findall(education_text)
        
        # Find GPAs
        gpa_matches = self.patterns['gpa'].findall(education_text)
        
        # Create education objects
        for i, match in enumerate(degree_matches):
            education = Education()
            education.degree = match.group().strip()
            
            # Try to match with nearby organizations
            if organizations:
                # Find organization closest to degree mention
                degree_pos = match.start()
                min_distance = float('inf')
                closest_org = None
                
                for org in organizations:
                    org_pos = education_text.find(org)
                    if org_pos != -1:
                        distance = abs(org_pos - degree_pos)
                        if distance < min_distance:
                            min_distance = distance
                            closest_org = org
                
                if closest_org:
                    education.institution = closest_org
            
            # Assign year if available
            if i < len(years):
                education.year = years[i]
            
            # Assign GPA if available
            if i < len(gpa_matches):
                education.gpa = gpa_matches[i]
            
            educations.append(education)
        
        # If no formal degree patterns found, look for university/college mentions
        if not educations:
            university_pattern = re.compile(r'(?:university|college|institute|school)\s+of\s+([^,\n]+)', re.IGNORECASE)
            uni_matches = university_pattern.finditer(text)
            
            for match in uni_matches:
                education = Education()
                education.institution = match.group().strip()
                if years:
                    education.year = years[0]
                educations.append(education)
                break  # Just take the first one
        
        return educations
    
    def extract_experience(self, text: str) -> List[Experience]:
        """Extract work experience information"""
        experiences = []
        
        # Common job title keywords for better matching
        job_title_keywords = [
            'Software Engineer', 'Data Scientist', 'Product Manager', 'Business Analyst',
            'Project Manager', 'Marketing Manager', 'Sales Representative', 'Consultant',
            'Developer', 'Analyst', 'Specialist', 'Coordinator', 'Director', 'Manager',
            'Lead', 'Senior', 'Junior', 'Intern', 'Executive', 'Associate', 'Principal',
            'Architect', 'Designer', 'Administrator', 'Technician', 'Supervisor'
        ]
        
        # Split by common experience section headers
        experience_sections = re.split(
            r'\b(?:experience|employment|work\s+history|career|professional\s+experience)\b', 
            text, 
            flags=re.IGNORECASE
        )
        
        if len(experience_sections) > 1:
            experience_text = experience_sections[1][:3000]  # Limit size
        else:
            experience_text = text
        
        # Use NLP to find organizations
        doc = self.nlp(experience_text)
        organizations = [ent.text for ent in doc.ents if ent.label_ == "ORG" and len(ent.text) > 1]
        
        # Find date ranges (experience duration)
        date_pattern = re.compile(r'(\d{4})\s*[-–]\s*(\d{4}|present|current)', re.IGNORECASE)
        date_matches = date_pattern.findall(experience_text)
        
        # Extract job titles and create experience objects
        lines = experience_text.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line or len(line) < 5:
                continue
            
            # Check if line contains job title keywords
            found_title = None
            for title in job_title_keywords:
                if title.lower() in line.lower():
                    found_title = title
                    break
            
            if found_title:
                experience = Experience()
                # Use the actual line as job title (may contain more specific info)
                experience.job_title = line[:150]  # Limit length
                
                # Find associated company
                # Look in current line and next few lines
                context_lines = lines[i:i+3] if i+3 < len(lines) else lines[i:]
                context_text = ' '.join(context_lines)
                
                for org in organizations:
                    if org.lower() in context_text.lower():
                        experience.company = org
                        break
                
                # Find duration
                if date_matches:
                    for start_year, end_year in date_matches:
                        experience.duration = f"{start_year} - {end_year}"
                        date_matches.remove((start_year, end_year))  # Use each date range only once
                        break
                
                # Add description from subsequent lines
                description_lines = []
                for j in range(i+1, min(i+5, len(lines))):
                    next_line = lines[j].strip()
                    if next_line and not any(keyword.lower() in next_line.lower() for keyword in job_title_keywords):
                        if not any(org.lower() in next_line.lower() for org in organizations):
                            description_lines.append(next_line)
                        else:
                            break
                    elif any(keyword.lower() in next_line.lower() for keyword in job_title_keywords):
                        break
                
                if description_lines:
                    experience.description = ' '.join(description_lines)[:300]  # Limit description
                
                experiences.append(experience)
        
        # Remove duplicates based on job title similarity
        unique_experiences = []
        for exp in experiences:
            is_duplicate = False
            for unique_exp in unique_experiences:
                if (exp.job_title and unique_exp.job_title and 
                    exp.job_title.lower()[:20] == unique_exp.job_title.lower()[:20]):
                    is_duplicate = True
                    break
            if not is_duplicate:
                unique_experiences.append(exp)
        
        return unique_experiences[:10]  # Limit to 10 most recent experiences
    
    def extract_certifications(self, text: str) -> List[str]:
        """Extract certifications and licenses"""
        certifications = set()
        
        # Common certification patterns
        cert_patterns = [
            r'\b(?:certified?|certification)\s+([^,.\n]{3,50})',
            r'\b([A-Z]{2,}\s+(?:certified?|certification)[^,.\n]{0,30})',
            r'\b(AWS|Google|Microsoft|Oracle|Cisco|CompTIA|PMI|Salesforce|Adobe)\s+[^,.\n]{3,50}',
            r'\b(?:licensed?|license)\s+([^,.\n]{3,50})',
        ]
        
        for pattern in cert_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                cert = match.group(1) if match.lastindex and match.group(1) else match.group()
                cert = cert.strip()
                if len(cert) > 3 and len(cert) < 100:  # Reasonable length
                    certifications.add(cert)
        
        # Look for common certification abbreviations
        cert_abbreviations = [
            'PMP', 'CISSP', 'CISA', 'CISM', 'CompTIA A+', 'Security+', 'Network+',
            'MCSE', 'MCSA', 'CCNA', 'CCNP', 'CCIE', 'AWS Solutions Architect',
            'Google Cloud Professional', 'Azure Fundamentals', 'Scrum Master',
            'Six Sigma', 'ITIL', 'CPA', 'CFA', 'FRM'
        ]
        
        text_upper = text.upper()
        for abbrev in cert_abbreviations:
            if abbrev.upper() in text_upper:
                certifications.add(abbrev)
        
        return sorted(list(certifications))
    
    def extract_languages(self, text: str) -> List[str]:
        """Extract languages from resume"""
        languages = set()
        
        # Common languages
        language_list = [
            'English', 'Spanish', 'French', 'German', 'Italian', 'Portuguese', 'Russian',
            'Chinese', 'Mandarin', 'Japanese', 'Korean', 'Arabic', 'Hindi', 'Dutch',
            'Swedish', 'Norwegian', 'Danish', 'Finnish', 'Polish', 'Greek', 'Turkish'
        ]
        
        # Look for language patterns
        language_patterns = [
            r'(?:languages?|linguistic|fluent|native|proficient):\s*([^.\n]{10,100})',
            r'(?:speak|spoken|speaking)\s+([^,.\n]{3,30})',
            r'(?:fluent|native|proficient)\s+(?:in\s+)?([A-Za-z]+)',
        ]
        
        for pattern in language_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                lang_text = match.group(1).lower()
                for lang in language_list:
                    if lang.lower() in lang_text:
                        languages.add(lang)
        
        # Direct matching
        text_lower = text.lower()
        for lang in language_list:
            if lang.lower() in text_lower:
                languages.add(lang)
        
        return sorted(list(languages))

class ResumeParser:
    """Main resume parser class with multi-threading support"""
    
    def __init__(self, max_workers: int = 4):
        """Initialize parser with thread pool"""
        self.max_workers = max_workers
        self.text_extractor = TextExtractor()
        self.entity_extractor = EntityExtractor()
        self.lock = threading.Lock()
        logger.info(f"Initialized Sereena Parser with {max_workers} workers")
    
    def parse_single_resume(self, file_path: str) -> Tuple[str, Optional[ResumeData]]:
        """Parse a single resume file"""
        start_time = time.time()
        
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Check file size (avoid very large files)
            file_size = file_path.stat().st_size
            if file_size > 50 * 1024 * 1024:  # 50MB limit
                raise ValueError(f"File too large: {file_size} bytes")
            
            # Extract text based on file extension
            if file_path.suffix.lower() == '.pdf':
                raw_text = self.text_extractor.extract_from_pdf(str(file_path))
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                raw_text = self.text_extractor.extract_from_docx(str(file_path))
            else:
                raise ValueError(f"Unsupported file format: {file_path.suffix}")
            
            if not raw_text or len(raw_text.strip()) < 50:
                raise ValueError("Insufficient text extracted from file")
            
            # Extract entities
            personal_info = self.entity_extractor.extract_personal_info(raw_text)
            skills = self.entity_extractor.extract_skills(raw_text)
            education = self.entity_extractor.extract_education(raw_text)
            experience = self.entity_extractor.extract_experience(raw_text)
            certifications = self.entity_extractor.extract_certifications(raw_text)
            languages = self.entity_extractor.extract_languages(raw_text)
            
            # Create summary (first meaningful sentences)
            sentences = re.split(r'[.!?]+', raw_text)
            meaningful_sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
            summary = '. '.join(meaningful_sentences[:3]) + '.' if meaningful_sentences else raw_text[:200]
            
            # Create resume data object
            resume_data = ResumeData(
                personal_info=personal_info,
                skills=skills,
                education=education,
                experience=experience,
                certifications=certifications,
                languages=languages,
                summary=summary[:500],  # Limit summary length
                raw_text=raw_text[:5000] if len(raw_text) > 5000 else raw_text  # Limit raw text storage
            )
            
            processing_time = time.time() - start_time
            logger.info(f"Successfully parsed {file_path.name} in {processing_time:.2f}s")
            
            return str(file_path), resume_data
            
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {str(e)}")
            return str(file_path), None
    
    def parse_multiple_resumes(self, file_paths: List[str]) -> Dict[str, Optional[ResumeData]]:
        """Parse multiple resumes using multi-threading"""
        if not file_paths:
            return {}
        
        results = {}
        successful = 0
        failed = 0
        
        logger.info(f"Starting batch processing of {len(file_paths)} files")
        
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            # Submit all tasks
            future_to_file = {
                executor.submit(self.parse_single_resume, file_path): file_path 
                for file_path in file_paths
            }
            
            # Collect results as they complete
            for future in as_completed(future_to_file):
                original_file_path = future_to_file[future]
                try:
                    file_key, resume_data = future.result()
                    with self.lock:
                        results[file_key] = resume_data
                        if resume_data:
                            successful += 1
                        else:
                            failed += 1
                except Exception as e:
                    logger.error(f"Thread error for {original_file_path}: {str(e)}")
                    with self.lock:
                        results[original_file_path] = None
                        failed += 1
        
        logger.info(f"Batch processing complete: {successful} successful, {failed} failed")
        return results
    
    def generate_insights(self, resume_data: ResumeData) -> Dict[str, Any]:
        """Generate insights from parsed resume data"""
        if not resume_data:
            return {}
        
        insights = {
            'profile_completeness': self._calculate_completeness(resume_data),
            'skill_categories': self._categorize_skills(resume_data.skills),
            'experience_summary': self._summarize_experience(resume_data.experience),
            'education_level': self._determine_education_level(resume_data.education),
            'career_progression': self._analyze_career_progression(resume_data.experience),
            'technical_score': self._calculate_technical_score(resume_data.skills),
            'language_diversity': len(resume_data.languages),
            'certification_count': len(resume_data.certifications),
        }
        
        return insights
    
    def _calculate_completeness(self, resume_data: ResumeData) -> float:
        """Calculate profile completeness percentage"""
        total_fields = 8
        completed_fields = 0
        
        if resume_data.personal_info.name:
            completed_fields += 1
        if resume_data.personal_info.email:
            completed_fields += 1
        if resume_data.personal_info.phone:
            completed_fields += 0.5
        if resume_data.skills:
            completed_fields += 1
        if resume_data.education:
            completed_fields += 1
        if resume_data.experience:
            completed_fields += 2  # Experience is more important
        if resume_data.certifications:
            completed_fields += 0.5
        if resume_data.languages:
            completed_fields += 0.5
        
        return min((completed_fields / total_fields) * 100, 100)
    
    def _categorize_skills(self, skills: List[str]) -> Dict[str, List[str]]:
        """Categorize skills into different categories"""
        categorized = {category: [] for category in self.entity_extractor.skills_db.keys()}
        uncategorized = []
        
        for skill in skills:
            categorized_flag = False
            for category, category_skills in self.entity_extractor.skills_db.items():
                if skill in category_skills:
                    categorized[category].append(skill)
                    categorized_flag = True
                    break
            
            if not categorized_flag:
                uncategorized.append(skill)
        
        # Remove empty categories
        categorized = {k: v for k, v in categorized.items() if v}
        
        if uncategorized:
            categorized['other'] = uncategorized
        
        return categorized
    
    def _summarize_experience(self, experiences: List[Experience]) -> Dict[str, Any]:
        """Summarize work experience"""
        companies = [exp.company for exp in experiences if exp.company]
        job_titles = [exp.job_title for exp in experiences if exp.job_title]
        
        # Extract years of experience from durations
        total_years = 0
        for exp in experiences:
            if exp.duration:
                # Try to extract year range
                year_match = re.findall(r'(\d{4})', exp.duration)
                if len(year_match) >= 2:
                    try:
                        start_year = int(year_match[0])
                        end_year = int(year_match[-1])
                        if end_year > start_year:
                            total_years += (end_year - start_year)
                    except ValueError:
                        pass
        
        return {
            'total_positions': len(experiences),
            'companies': list(set(companies)),
            'job_titles': list(set(job_titles)),
            'estimated_years': max(total_years, len(experiences))  # At least 1 year per position
        }
    
    def _determine_education_level(self, educations: List[Education]) -> str:
        """Determine highest education level"""
        if not educations:
            return "Not specified"
        
        levels = {
            'phd': 5, 'doctorate': 5, 'doctoral': 5, 'ph.d': 5,
            'master': 4, 'mba': 4, 'ms': 4, 'ma': 4, 'm.s': 4, 'm.a': 4,
            'bachelor': 3, 'bs': 3, 'ba': 3, 'b.s': 3, 'b.a': 3,
            'associate': 2, 'diploma': 1, 'certificate': 1
        }
        
        highest_level = 0
        highest_degree = "Not specified"
        
        for education in educations:
            if education.degree:
                degree_lower = education.degree.lower()
                for level_name, level_value in levels.items():
                    if level_name in degree_lower and level_value > highest_level:
                        highest_level = level_value
                        highest_degree = education.degree
        
        return highest_degree if highest_degree != "Not specified" else "Bachelor's Level"
    
    def _analyze_career_progression(self, experiences: List[Experience]) -> str:
        """Analyze career progression pattern"""
        if len(experiences) < 2:
            return "Insufficient data for analysis"
        
        # Keywords that indicate seniority levels
        seniority_keywords = {
            'intern': 1, 'trainee': 1,
            'junior': 2, 'associate': 2, 'assistant': 2,
            'mid': 3, 'regular': 3,
            'senior': 4, 'lead': 5, 'principal': 6,
            'manager': 7, 'director': 8, 'vp': 9, 'president': 10, 'ceo': 10
        }
        
        progression_scores = []
        for exp in experiences:
            if exp.job_title:
                title_lower = exp.job_title.lower()
                score = 3  # Default middle-level score
                for keyword, value in seniority_keywords.items():
                    if keyword in title_lower:
                        score = value
                        break
                progression_scores.append(score)
        
        if len(progression_scores) >= 2:
            # Compare first and last positions
            if progression_scores[-1] > progression_scores[0]:
                return "Upward career progression"
            elif progression_scores[-1] < progression_scores[0]:
                return "Career transition or step back"
            else:
                return "Lateral career movement"
        
        return "Stable career path"
    
    def _calculate_technical_score(self, skills: List[str]) -> int:
        """Calculate technical proficiency score out of 100"""
        if not skills:
            return 0
        
        # Technical skill categories with different weights
        technical_categories = {
            'programming': 3,
            'databases': 2,
            'cloud_platforms': 2,
            'tools': 1,
            'data_science': 3,
            'web_technologies': 2,
            'mobile': 2,
            'devops': 2
        }
        
        total_score = 0
        for skill in skills:
            for category, category_skills in self.entity_extractor.skills_db.items():
                if skill in category_skills and category in technical_categories:
                    total_score += technical_categories[category]
                    break
        
        # Normalize to 0-100 scale
        return min(total_score * 2, 100)
    
    def export_results(self, results: Dict[str, Optional[ResumeData]], output_file: str = "parsed_resumes.json"):
        """Export results to JSON file with insights"""
        export_data = {}
        
        for file_path, resume_data in results.items():
            if resume_data:
                export_data[file_path] = {
                    'resume_data': asdict(resume_data),
                    'insights': self.generate_insights(resume_data),
                    'parsing_timestamp': time.strftime('%Y-%m-%d %H:%M:%S')
                }
            else:
                export_data[file_path] = {
                    'error': 'Failed to parse',
                    'parsing_timestamp': time.strftime('%Y-%m-%d %H:%M:%S')
                }
        
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False, default=str)
            
            logger.info(f"Results exported to {output_file}")
            
            # Print summary statistics
            successful_parses = sum(1 for v in results.values() if v is not None)
            total_files = len(results)
            
            print(f"\n📊 Parsing Summary:")
            print(f"   Total files processed: {total_files}")
            print(f"   Successfully parsed: {successful_parses}")
            print(f"   Failed to parse: {total_files - successful_parses}")
            print(f"   Success rate: {(successful_parses/total_files)*100:.1f}%")
            
        except Exception as e:
            logger.error(f"Failed to export results: {str(e)}")
            raise
    
    def get_parsing_statistics(self, results: Dict[str, Optional[ResumeData]]) -> Dict[str, Any]:
        """Generate parsing statistics"""
        successful_results = [data for data in results.values() if data is not None]
        
        if not successful_results:
            return {"message": "No successful parses to analyze"}
        
        # Aggregate statistics
        total_skills = sum(len(data.skills) for data in successful_results)
        total_experiences = sum(len(data.experience) for data in successful_results)
        total_educations = sum(len(data.education) for data in successful_results)
        total_certifications = sum(len(data.certifications) for data in successful_results)
        
        # Most common skills
        skill_counter = {}
        for data in successful_results:
            for skill in data.skills:
                skill_counter[skill] = skill_counter.get(skill, 0) + 1
        
        top_skills = sorted(skill_counter.items(), key=lambda x: x[1], reverse=True)[:10]
        
        # Most common companies
        company_counter = {}
        for data in successful_results:
            for exp in data.experience:
                if exp.company:
                    company_counter[exp.company] = company_counter.get(exp.company, 0) + 1
        
        top_companies = sorted(company_counter.items(), key=lambda x: x[1], reverse=True)[:10]
        
        return {
            'total_resumes_parsed': len(successful_results),
            'average_skills_per_resume': round(total_skills / len(successful_results), 1),
            'average_experience_per_resume': round(total_experiences / len(successful_results), 1),
            'average_education_per_resume': round(total_educations / len(successful_results), 1),
            'average_certifications_per_resume': round(total_certifications / len(successful_results), 1),
            'top_skills': [{"skill": skill, "count": count} for skill, count in top_skills],
            'top_companies': [{"company": company, "count": count} for company, count in top_companies],
        }

def print_resume_summary(resume_data: ResumeData, insights: Dict[str, Any]):
    """Print a formatted summary of parsed resume data"""
    print("\n" + "="*60)
    print("📋 RESUME PARSING RESULTS")
    print("="*60)
    
    # Personal Information
    print("\n👤 PERSONAL INFORMATION:")
    print(f"   Name: {resume_data.personal_info.name or 'Not found'}")
    print(f"   Email: {resume_data.personal_info.email or 'Not found'}")
    print(f"   Phone: {resume_data.personal_info.phone or 'Not found'}")
    print(f"   Location: {resume_data.personal_info.location or 'Not found'}")
    
    # Skills
    print(f"\n🛠️  SKILLS ({len(resume_data.skills)} found):")
    if resume_data.skills:
        for i, skill in enumerate(resume_data.skills[:15], 1):  # Show first 15 skills
            print(f"   {i:2}. {skill}")
        if len(resume_data.skills) > 15:
            print(f"   ... and {len(resume_data.skills) - 15} more")
    else:
        print("   No skills found")
    
    # Experience
    print(f"\n💼 WORK EXPERIENCE ({len(resume_data.experience)} positions):")
    for i, exp in enumerate(resume_data.experience, 1):
        print(f"   {i}. {exp.job_title or 'Unknown Position'}")
        if exp.company:
            print(f"      Company: {exp.company}")
        if exp.duration:
            print(f"      Duration: {exp.duration}")
    
    # Education
    print(f"\n🎓 EDUCATION ({len(resume_data.education)} entries):")
    for i, edu in enumerate(resume_data.education, 1):
        print(f"   {i}. {edu.degree or 'Unknown Degree'}")
        if edu.institution:
            print(f"      Institution: {edu.institution}")
        if edu.year:
            print(f"      Year: {edu.year}")
    
    # Certifications
    if resume_data.certifications:
        print(f"\n🏆 CERTIFICATIONS ({len(resume_data.certifications)}):")
        for i, cert in enumerate(resume_data.certifications, 1):
            print(f"   {i}. {cert}")
    
    # Languages
    if resume_data.languages:
        print(f"\n🌍 LANGUAGES ({len(resume_data.languages)}):")
        print(f"   {', '.join(resume_data.languages)}")
    
    # Insights
    print(f"\n📊 INSIGHTS:")
    print(f"   Profile Completeness: {insights.get('profile_completeness', 0):.1f}%")
    print(f"   Technical Score: {insights.get('technical_score', 0)}/100")
    print(f"   Education Level: {insights.get('education_level', 'Unknown')}")
    print(f"   Career Progression: {insights.get('career_progression', 'Unknown')}")
    
    print("\n" + "="*60)

def cli_main():
    """Command line interface for Sereena Parser"""
    parser = argparse.ArgumentParser(
        description='Sereena Parser - Advanced Resume Parsing Tool',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  sereena-parser resume.pdf                          # Parse single file
  sereena-parser *.pdf *.docx                       # Parse multiple files
  sereena-parser resume.pdf -o results.json         # Custom output file
  sereena-parser *.pdf -w 8 --verbose               # 8 workers, verbose output
  sereena-parser resume.pdf --summary               # Show detailed summary
        """
    )
    
    parser.add_argument(
        'files', 
        nargs='+', 
        help='Resume files to parse (PDF or DOCX format)'
    )
    parser.add_argument(
        '-o', '--output', 
        default='parsed_resumes.json', 
        help='Output JSON file (default: parsed_resumes.json)'
    )
    parser.add_argument(
        '-w', '--workers', 
        type=int, 
        default=4, 
        help='Number of worker threads (default: 4)'
    )
    parser.add_argument(
        '-v', '--verbose', 
        action='store_true', 
        help='Enable verbose logging'
    )
    parser.add_argument(
        '--summary', 
        action='store_true', 
        help='Show detailed summary for single file parsing'
    )
    parser.add_argument(
        '--stats', 
        action='store_true', 
        help='Show parsing statistics for batch processing'
    )
    parser.add_argument(
        '--version', 
        action='version', 
        version='Sereena Parser 1.0.0'
    )
    
    args = parser.parse_args()
    
    # Configure logging
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
        logging.getLogger('sereena_parser').setLevel(logging.DEBUG)
    
    # Validate input files
    valid_files = []
    for file_path in args.files:
        path = Path(file_path)
        if path.exists() and path.suffix.lower() in ['.pdf', '.docx', '.doc']:
            valid_files.append(str(path))
        else:
            print(f"⚠️  Warning: Skipping invalid or non-existent file: {file_path}")
    
    if not valid_files:
        print("❌ Error: No valid files to process")
        sys.exit(1)
    
    print(f"🚀 Starting Sereena Parser...")
    print(f"📁 Files to process: {len(valid_files)}")
    print(f"👥 Worker threads: {args.workers}")
    
    # Initialize parser
    resume_parser = ResumeParser(max_workers=args.workers)
    
    try:
        # Parse files
        if len(valid_files) == 1 and args.summary:
            # Single file with detailed summary
            file_path, resume_data = resume_parser.parse_single_resume(valid_files[0])
            
            if resume_data:
                insights = resume_parser.generate_insights(resume_data)
                print_resume_summary(resume_data, insights)
                
                # Also save to JSON
                results = {file_path: resume_data}
                resume_parser.export_results(results, args.output)
            else:
                print(f"❌ Failed to parse: {valid_files[0]}")
                sys.exit(1)
        
        else:
            # Batch processing
            results = resume_parser.parse_multiple_resumes(valid_files)
            
            # Export results
            resume_parser.export_results(results, args.output)
            
            # Show statistics if requested
            if args.stats:
                stats = resume_parser.get_parsing_statistics(results)
                print(f"\n📈 PARSING STATISTICS:")
                print(f"   Total resumes parsed: {stats.get('total_resumes_parsed', 0)}")
                print(f"   Average skills per resume: {stats.get('average_skills_per_resume', 0)}")
                print(f"   Average experience per resume: {stats.get('average_experience_per_resume', 0)}")
                
                if 'top_skills' in stats and stats['top_skills']:
                    print(f"\n🔝 TOP SKILLS:")
                    for i, skill_data in enumerate(stats['top_skills'][:5], 1):
                        print(f"   {i}. {skill_data['skill']} ({skill_data['count']} resumes)")
        
        print(f"\n✅ Processing complete! Results saved to: {args.output}")
        
    except KeyboardInterrupt:
        print(f"\n⚠️  Processing interrupted by user")
        sys.exit(1)
    except Exception as e:
        logger.error(f"Fatal error: {str(e)}")
        print(f"❌ Fatal error: {str(e)}")
        sys.exit(1)

def main():
    """Entry point for the package"""
    cli_main()

if __name__ == "__main__":
    cli_main()