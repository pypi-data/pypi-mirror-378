from docx_asn1 import extract_text_from_docx
from docx import Document
from pathlib import Path


def test_decode():
    path_to_doc = Path(__file__).parent / "test.docx"
    doc = Document()
    doc.add_paragraph("test1")
    doc.add_paragraph("-- ASN1START")
    doc.add_paragraph("\ttest2 ")
    doc.add_paragraph("-- ASN1STOP")
    doc.add_paragraph("test3")
    doc.save(path_to_doc)
    res = extract_text_from_docx(path_to_doc)
    assert res == "\ttest2 "


def test_decode_file_smarttags():
    path_to_doc = Path(__file__).parent / "36413-ge0.docx"
    path_to_txt = Path(__file__).parent / "36413-ge0.asn"
    res = extract_text_from_docx(
        path_to_doc, start_text="-- ***************", stop_text="END", add=True
    )
    with open(path_to_txt, "w", encoding="utf-8") as file:
        file.write(res)
    # should handle smarttags
    # assert "CSGStatus ::=" in res
    assert "CSGMembershipStatus ::=" in res


if __name__ == "__main__":
    test_decode_file_smarttags()
