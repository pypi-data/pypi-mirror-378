# encoding: utf-8

import random
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches

from .excel_util import excel_read
from .file_util import FileUtil
from .perf_util import PerfUtil


def func_report(docx, xlsx, sheet, project_num, path):
    data, max_column = excel_read(xlsx, sheet)
    columns = max_column - 3

    for single in data:
        docx_input = Document(rf'{docx}')
        title = single[0]
        # print(title)
        # 正则提取标题名称
        pattern = re.compile(r'^\d*\.').search(title)
        if pattern is not None:
            ttl = title[pattern.span()[1]:]
        else:
            ttl = title
        # with open(r"../config/name.txt", 'w') as f:
        #     f.write(ttl)
        # conclusion = get_perf_report(title=ttl)
        year = single[1]
        month = single[2]
        day_before = random.randint(1, 12)
        day_after = None
        if project_num != 4:
            day_after = random.randint(13, 28)
        if project_num == 4:
            day_after = day_before + 3
        # 动态赋值
        for k in range(1, columns + 1):
            exec(f'param{k} = "{single[2 + k]}"', globals())
        params = globals()

        docx_input.styles['Normal'].font.name = u'宋体'
        # noinspection PyProtectedMember
        docx_input.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        # 这里修改一下聚合报告图片
        perf_util = PerfUtil(title)
        perf_util.change_summary()
        perf_util.change_threads()
        conclusion = perf_util.get_conclusion()

        for i in docx_input.paragraphs:
            for j in i.runs:
                if "ttl" in j.text:
                    j.text = j.text.replace("ttl", ttl)
                if "title" in j.text:
                    j.text = j.text.replace("title", title)
                if "year" in j.text:
                    j.text = j.text.replace("year", str(year))
                if "month" in j.text:
                    j.text = j.text.replace("month", str(month))
                if "day_before" in j.text:
                    j.text = j.text.replace("day_before", str(day_before))
                if "day_after" in j.text:
                    j.text = j.text.replace("day_after", str(day_after))
                if "summary" in j.text:
                    j.text = j.text.replace('summary', '')
                    j.add_picture("./assets/summary.png", width=Inches(6.2))
                if "threads" in j.text:
                    j.text = j.text.replace('threads', '')
                    j.add_picture("./assets/threads.png", width=Inches(6.2))
                if "conclusion" in j.text:
                    j.text = j.text.replace("conclusion", str(conclusion))

        for table in docx_input.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "title" in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.text = run.text.replace("title", title)
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    if "ttl" in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.text = run.text.replace("ttl", ttl)
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    for k in range(1, columns + 1):
                        # 利用命名空间获取动态变量的值
                        param = params.get('param' + str(k))
                        if f"param{k}" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace(f"param{k}", param)
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        save_path = FileUtil(project_num=project_num, path=path, year=year, month=month, title=title, type=1).folder_path()
        docx_input.save(save_path)

        print(f"Save success! title:{title}")
