# encoding: utf-8

import re
import random
import string
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from .excel_util import excel_read
from .file_util import FileUtil


def safe_report(docx, xlsx, sheet, project_num, path):
    data, max_column = excel_read(xlsx, sheet)

    for single in data:
        docx_input = Document(rf'{docx}')
        title = single[0]
        # 一码通专用
        param1 = single[3]
        # 正则提取标题名称
        pattern = re.compile(r'^\d*\.').search(title)
        if pattern is not None:
            ttl = title[pattern.span()[1]:]
        else:
            ttl = title
        year = single[1]
        month = single[2]
        day_before = random.randint(8, 12)
        day_after = random.randint(9, 17)

        if project_num != 4:
            error1 = random.randint(0, 1)
            error2 = random.randint(0, 1)
            error3 = random.randint(0, 1)
            error4 = random.randint(0, 1)
            error5 = random.randint(0, 1)
            error6 = random.randint(0, 1)
            error = error1 + error2 + error3 + error4 + error5 + error6
            rand1 = random.choice(string.ascii_letters)
            rand2 = random.choice(string.ascii_letters)
            docxurl = "http://220.196.242.107/smartGuide/" + str(rand1) + str(rand2)

            safety_content = ''
            if error1 != 0:
                safety_content = safety_content + '修改服务器配置，TLS协议仅支持TLSv1.2及以上。\n'
            if error2 != 0:
                safety_content = safety_content + '修改服务端程序，给 HTTP 响应头加上 X-Permitted-Cross-Domain-Policies。\n'
            if error3 != 0:
                safety_content = safety_content + '建议给cookie加上HTTPOnly属性。\n'
            if error4 != 0:
                safety_content = safety_content + '对cookie设置安全标志。\n'
            if error5 != 0:
                safety_content = safety_content + '做如下安全加固：\n1.必需尽快找到无效链接网页，修改或删除。\n2.利用Robost.txt文件屏蔽，不让搜索抓取。\n'
            if error6 != 0:
                safety_content = safety_content + '升级到最新版本。\n'

            docx_input.styles['Normal'].font.name = u'宋体'
            # noinspection PyProtectedMember
            docx_input.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

            for i in docx_input.paragraphs:
                for j in i.runs:
                    if "title" in j.text:
                        j.text = j.text.replace("title", title)
                    if "ttl" in j.text:
                        j.text = j.text.replace("ttl", ttl)
                    if "year" in j.text:
                        j.text = j.text.replace("year", str(year))
                    if "month" in j.text:
                        j.text = j.text.replace("month", str(month))
                    if "day_before" in j.text:
                        j.text = j.text.replace("day_before", str(day_before))
                    if "day_after" in j.text:
                        j.text = j.text.replace("day_after", str(day_after))
                    if "error" in j.text:
                        j.text = j.text.replace("error", str(error))
                    if "safety_content" in j.text:
                        j.text = j.text.replace("safety_content", safety_content)

            for table in docx_input.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "title" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("title", title)
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        # 一码通专用
                        if "param1" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("param1", param1)
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        if "ttl" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("ttl", ttl)
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        if "error1" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("error1", str(error1))
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if "error2" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("error2", str(error2))
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if "error3" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("error3", str(error3))
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if "error4" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("error4", str(error4))
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if "error5" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("error5", str(error5))
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if "error6" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("error6", str(error6))
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if "error" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("error", str(error))
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if "<url>" in cell.text:
                            # print(True)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    # print(run.text)
                                    run.text = run.text.replace("<url>", str(docxurl))
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if project_num == 4:
            error1 = random.randint(0, 1)
            error2 = random.randint(0, 1)
            error3 = random.randint(0, 1)
            error = error1 + error2 + error3
            rand1 = random.choice(string.ascii_letters)
            rand2 = random.choice(string.ascii_letters)
            docxurl = "http://220.196.242.107/smartGuide/" + str(rand1) + str(rand2)

            docx_input.styles['Normal'].font.name = u'宋体'
            # noinspection PyProtectedMember
            docx_input.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

            for i in docx_input.paragraphs:
                for j in i.runs:
                    if "title" in j.text:
                        j.text = j.text.replace("title", title)
                    if "ttl" in j.text:
                        j.text = j.text.replace("ttl", ttl)
                    if "year" in j.text:
                        j.text = j.text.replace("year", str(year))
                    if "month" in j.text:
                        j.text = j.text.replace("month", str(month))
                    if "day_before" in j.text:
                        j.text = j.text.replace("day_before", str(day_before))
                    if "day_after" in j.text:
                        j.text = j.text.replace("day_after", str(day_after))
                    if "error" in j.text:
                        j.text = j.text.replace("error", str(error))

            for table in docx_input.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if "title" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("title", title)
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        # 一码通专用
                        if "param1" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("param1", param1)
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        if "ttl" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("ttl", ttl)
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        if "err_1" in cell.text:
                            # print(True)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    # print(run.text)
                                    run.text = run.text.replace("err_1", str(error1))
                                    # print(run.text)
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            print(str(error1))
                        if "err_2" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("err_2", str(error2))
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if "err_3" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("err_3", str(error3))
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if "error" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("error", str(error))
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        if "<url>" in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.text = run.text.replace("<url>", str(docxurl))
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        save_path = FileUtil(project_num=project_num, path=path, year=year, month=month, title=title, type=2).folder_path()
        docx_input.save(save_path)

        print(f"Save success! title:{title}")
